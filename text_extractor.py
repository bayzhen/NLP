# text_extractor.py
import pdfplumber
import openpyxl
import win32com.client as win32
from typing import Any, Dict
from contextlib import contextmanager
from config import Config
from pathlib import Path
from datetime import datetime
import xlrd
import os
from new_docx_handler import NewDocxHandler
from multiprocessing import Process, Queue
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
import time


class FileTooLargeError(Exception):
    """文件处理超时异常"""
    pass


class TimeoutGuard:
    """超时保护上下文管理器"""

    def __init__(self, timeout):
        self.timeout = timeout
        self.start_time = None

    def __enter__(self):
        self.start_time = time.time()
        return self

    def __exit__(self, *args):
        pass

    def check_timeout(self):
        duration = time.time() - self.start_time
        if duration > self.timeout:
            raise FileTooLargeError(f"处理超时，超过 {self.timeout} 秒")
        else:
            print(f"时间已过：{duration:2f}")


class TextExtractor:
    """多格式文本提取器（新增超时机制）"""

    @classmethod
    def extract(cls, file_path: str) -> str:
        """统一入口方法（添加超时控制）"""
        ext = Path(file_path).suffix.lower()
        handler = getattr(cls, f"_handle_{ext[1:]}", cls._handle_unsupported)
        return handler(file_path)

    @staticmethod
    def _handle_unsupported(file_path: str) -> str:
        print("不支持的文件：", file_path)
        return ""

    @staticmethod
    def _handle_txt(file_path: str) -> str:
        """分块读取文本文件"""
        content = []
        chunk_size = 4096  # 4KB为块单位
        with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer, \
                open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            while True:
                timer.check_timeout()
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                content.append(chunk)
        return "".join(content)

    @staticmethod
    def _handle_pdf(file_path: str) -> str:
        """分页处理PDF，支持超时中断"""
        text = []
        with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer, pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                timer.check_timeout()  # 每页前检查超时
                if text_content := page.extract_text():
                    text.append(text_content)
        return "\n".join(text)

    @classmethod
    def _handle_doc(cls, file_path: str) -> str:
        """改进的.doc处理，增加超时检查"""
        file_path = os.path.abspath(file_path)
        word = None
        try:
            with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer:
                word = win32.gencache.EnsureDispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(file_path)
                timer.check_timeout()  # 关键操作后检查

                # 分段落读取以便超时检查
                content = []
                for para in doc.Paragraphs:
                    timer.check_timeout()
                    content.append(para.Range.Text)
                return "\n".join(content)
        finally:
            if 'doc' in locals() and doc:
                doc.Close(SaveChanges=False)
            if word:
                word.Quit()

    @staticmethod
    def _handle_docx(file_path: str) -> str:
        """处理新版Word文档，提取所有文本内容（包括段落、表格、页眉、页脚）"""
        result = []
        try:
            doc = Document(file_path)
            with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer:
                # 处理正文中的段落和表格
                for block in NewDocxHandler.iter_block_items(doc):
                    timer.check_timeout()
                    if isinstance(block, Paragraph):
                        text = block.text.strip()
                        if text:
                            result.append(text)
                    elif isinstance(block, Table):
                        NewDocxHandler.process_table(block, result)

                # 处理页眉和页脚
                for section in doc.sections:
                    timer.check_timeout()
                    for header in [section.header, section.first_page_header]:
                        timer.check_timeout()
                        for para in header.paragraphs:
                            timer.check_timeout()
                            text = para.text.strip()
                            if text:
                                result.append(text)
                    for footer in [section.footer, section.first_page_footer]:
                        timer.check_timeout()
                        for para in footer.paragraphs:
                            timer.check_timeout()
                            text = para.text.strip()
                            if text:
                                result.append(text)
        except Exception as e:
            raise ValueError(f"无法处理Word文档: {e}")
        return "\n".join(result)

    @classmethod
    def _handle_xls(cls, file_path: str) -> str:
        """处理旧版Excel文件"""
        workbook = xlrd.open_workbook(file_path)
        all_text = ""
        with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer:
            for sheet_name in workbook.sheet_names():
                timer.check_timeout()
                sheet = workbook.sheet_by_name(sheet_name)
                for row_idx in range(sheet.nrows):
                    timer.check_timeout()
                    row_data = []
                    for col_idx in range(sheet.ncols):
                        timer.check_timeout()
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        # 将内容转为字符串（处理数字、日期等类型）
                        if isinstance(cell_value, float):
                            cell_value = str(int(cell_value)) if cell_value.is_integer() else str(cell_value)
                        else:
                            cell_value = str(cell_value)
                        row_data.append(cell_value)
                    for row in row_data:
                        timer.check_timeout()
                        all_text += "\t" + row
                    all_text += "\n"
            return all_text

    @classmethod
    def _handle_xlsx(cls, file_path: str) -> str:
        """处理新版Excel文件"""
        with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            xlsx_str = ""
            xlsx_dict = cls._process_excel(workbook)
            for key in xlsx_dict.keys():
                timer.check_timeout()
                value = xlsx_dict[key]
                xlsx_str += key + "\n"
                xlsx_str += value + "\n"
            return xlsx_str

    @classmethod
    def _process_excel(cls, workbook: Any) -> Dict[str, str]:
        """通用Excel处理逻辑"""
        results = {}
        with TimeoutGuard(Config.PROCESS_TIMEOUT) as timer:
            for sheet in workbook.worksheets:
                timer.check_timeout()
                lines = []
                for row in sheet.iter_rows():
                    timer.check_timeout()
                    cells = [cls._format_cell(cell) for cell in row]
                    if any(cells):
                        lines.append(Config.CELL_DELIMITER.join(cells))
                results[sheet.title] = Config.LINE_DELIMITER.join(lines)
            return results

    @staticmethod
    def _format_cell(cell: Any) -> str:
        """格式化单元格内容"""
        value = cell.value
        if value is None:
            return ""
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        return str(value).strip()
