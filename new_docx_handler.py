from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

class NewDocxHandler:
    @staticmethod
    def handle_docx(file_path: str) -> str:
        """处理新版Word文档，提取所有文本内容（包括段落、表格、页眉、页脚）"""
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"无法打开Word文档: {e}")

        result = []

        # 处理正文中的段落和表格
        for block in NewDocxHandler.iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    result.append(text)
            elif isinstance(block, Table):
                NewDocxHandler.process_table(block, result)

        # 处理页眉和页脚
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        result.append(text)
            for footer in [section.footer, section.first_page_footer]:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        result.append(text)

        return "\n".join(result)

    @staticmethod
    def iter_block_items(parent):
        """递归遍历文档中的所有块级元素（段落和表格）"""
        from docx.oxml.text.run import CT_R
        for elem in parent.element.body.iterchildren():
            if isinstance(elem, CT_P):
                yield Paragraph(elem, parent)
            elif isinstance(elem, CT_Tbl):
                table = Table(elem, parent)
                # 处理嵌套表格
                for row in table.rows:
                    for cell in row.cells:
                        for child_elem in cell._element.iterchildren():
                            if isinstance(child_elem, CT_P):
                                yield Paragraph(child_elem, parent)
                            elif isinstance(child_elem, CT_Tbl):
                                nested_table = Table(child_elem, parent)
                                yield nested_table
                            elif isinstance(child_elem, CT_R):
                                # 处理文本框内的文字（简单情况）
                                text = NewDocxHandler.extract_inline_text(child_elem)
                                if text:
                                    yield text
                yield table

    @staticmethod
    def process_table(table, result):
        """递归处理表格中的文本（包括嵌套表格）"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        result.append(text)
                # 处理嵌套表格
                for nested_table in cell.tables:
                    NewDocxHandler.process_table(nested_table, result)

    @staticmethod
    def extract_inline_text(run_element):
        """提取内联文本（如简单文本框内容）"""
        texts = []
        for t in run_element.itertext():
            texts.append(t.strip())
        return ' '.join(texts).strip()
